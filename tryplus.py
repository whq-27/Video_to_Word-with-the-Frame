import os
import sys
import shutil
import cv2
import re
import whisper
import yt_dlp
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import warnings

# --- 必须依赖 ---
try:
    from openai import OpenAI
except ImportError:
    print("❌ 缺少 openai 库，请运行: pip install openai")
    sys.exit(1)

warnings.filterwarnings("ignore")

# ================= 配置区域 =================
# 1. B站 Cookie
COOKIE_PATH = '/big-data/person/wanghaoqi/try/www.bilibili.com_cookies.txt'

# 2. 默认存储根目录
DEFAULT_OUTPUT_BASE = "/big-data/person/wanghaoqi/try/output"

# 3. LLM 配置
LLM_API_KEY = "sk-*************************************" 
LLM_BASE_URL = "https://api.deepseek.com"          
LLM_MODEL = "deepseek-chat"                       

# 4. 屏幕适配参数
MAX_SUBTITLE_LENGTH = 20  # 每行字幕最大字数
# ===========================================

class VideoToWordConverter:
    def __init__(self, model_size="base"):
        self.check_ffmpeg()
        self.model_size = model_size
        self.model = None
        self.video_path = None      
        self.img_output_dir = None
        self.video_title_stem = "video_report"
        
        self.llm_client = None
        if LLM_API_KEY:
            try:
                self.llm_client = OpenAI(api_key=LLM_API_KEY, base_url=LLM_BASE_URL)
                print(f"[系统] LLM 客户端已就绪: {LLM_MODEL}")
            except Exception as e:
                print(f"⚠️ LLM 初始化失败: {e}")

    def check_ffmpeg(self):
        if not shutil.which("ffmpeg"):
            print("❌ 错误: 未检测到 FFmpeg。")
            sys.exit(1)

    def _load_model(self):
        if self.model is None:
            print(f"[系统] 正在加载 Whisper 模型 ({self.model_size})...")
            try:
                self.model = whisper.load_model(self.model_size)
            except Exception as e:
                print(f"❌ 模型加载失败: {e}")
                sys.exit(1)

    def sanitize_filename(self, title):
        return re.sub(r'[\\/*?:"<>|]', "", title).strip()

    def prepare_source(self, user_input, output_dir):
        user_input = user_input.strip('"').strip("'")
        if os.path.exists(user_input):
            abs_path = os.path.abspath(user_input)
            print(f"✅ 检测到本地文件: {abs_path}")
            self.video_path = abs_path
            filename = os.path.basename(abs_path)
            self.video_title_stem = os.path.splitext(filename)[0]
            return True
        elif user_input.startswith(('http://', 'https://', 'www.')):
            print(f"🌐 检测到网络链接，准备获取标题并下载...")
            return self.download_video(user_input, output_dir)
        else:
            print("❌ 输入无效。")
            return False

    def download_progress_hook(self, d):
        if d['status'] == 'finished':
            print(f"    -> ✅ 流下载完成 ({d.get('_total_bytes_str', '未知')})...")

    def download_video(self, url, output_dir):
        use_cookie = os.path.exists(COOKIE_PATH)
        info_opts = {'quiet': True, 'no_warnings': True, 'user_agent': 'Mozilla/5.0...'}
        if use_cookie: info_opts['cookiefile'] = COOKIE_PATH

        print(f"[1/5] 获取视频标题...")
        try:
            with yt_dlp.YoutubeDL(info_opts) as ydl:
                info = ydl.extract_info(url, download=False)
                title = info.get('title', 'downloaded_video')
            self.video_title_stem = self.sanitize_filename(title)
            print(f"   -> 标题: {title}")
        except:
            self.video_title_stem = f"video_{datetime.now().strftime('%H%M%S')}"

        ydl_opts = {
            'format': 'bestvideo[vcodec^=avc]+bestaudio/best', 
            'merge_output_format': 'mp4',
            'outtmpl': os.path.join(output_dir, f"{self.video_title_stem}.%(ext)s"),
            'quiet': True, 'no_warnings': True,
            'progress_hooks': [self.download_progress_hook],
            'user_agent': 'Mozilla/5.0...'
        }
        if use_cookie: ydl_opts['cookiefile'] = COOKIE_PATH
        
        try:
            print(f"[2/5] 正在下载素材 (H.264)...")
            with yt_dlp.YoutubeDL(ydl_opts) as ydl:
                ydl.download([url])
            
            self.video_path = None
            for f in os.listdir(output_dir):
                if f == "images": continue
                if f.lower().endswith(('.mp4', '.mkv', '.webm')):
                    self.video_path = os.path.abspath(os.path.join(output_dir, f))
                    break
            if not self.video_path:
                self.video_path = os.path.abspath(os.path.join(output_dir, f"{self.video_title_stem}.mp4"))
            print(f"✅ 下载完成: {self.video_path}")
            return True
        except Exception as e:
            print(f"❌ 下载失败: {e}")
            return False

    # ================= 核心处理逻辑 =================

    def generate_summary(self, sentences_list):
        if not self.llm_client or not sentences_list: return None
        full_text = "".join([s['text'] for s in sentences_list])
        print(f"   -> 正在生成 AI 总结 (全文共 {len(full_text)} 字)...")
        prompt = """你是一个专业的会议纪要助手。请根据视频字幕生成一份结构化总结。
            要求：
            1. **一句话摘要**：精炼概括。
            2. **核心观点**：列出3-5个关键点。
            3. **详细脉络**：按逻辑梳理内容。
            请使用 Markdown 格式 (###, **, - )。"""
        try:
            safe_text = full_text[:30000] 
            response = self.llm_client.chat.completions.create(
                model=LLM_MODEL,
                messages=[
                    {"role": "system", "content": prompt},
                    {"role": "user", "content": f"字幕内容：\n{safe_text}"}
                ],
                stream=False
            )
            return response.choices[0].message.content
        except Exception as e:
            print(f"⚠️ 总结生成失败: {e}")
            return None

    def clean_text_minimal(self, text):
        if not self.llm_client or len(text) < 4: return text
        prompt = f"""任务：极简清洗字幕。
            原则：
            1. 仅删除口误、重复词、语气助词（呃、那个）。
            2. 【严禁】改写句子结构。
            3. 【严禁】删减实义内容。
            原文：{text}"""
        try:
            response = self.llm_client.chat.completions.create(
                model=LLM_MODEL,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.1, 
                stream=False
            )
            cleaned = response.choices[0].message.content.strip()
            if not cleaned: return text
            return cleaned
        except:
            return text

    def smart_split_subtitle(self, start, end, text, max_chars):
        total_duration = end - start
        pattern = r'([，。？！；：、,?!;:\s])'
        parts = re.split(pattern, text)
        parts = [p for p in parts if p]
        
        chunks = []
        current_chunk = ""
        
        for part in parts:
            if len(part) == 1 and re.match(pattern, part):
                current_chunk += part
                continue
            if len(current_chunk) + len(part) <= max_chars:
                current_chunk += part
            else:
                if current_chunk: chunks.append(current_chunk)
                current_chunk = ""
                if len(part) > max_chars:
                    for k in range(0, len(part), max_chars):
                        chunks.append(part[k:k+max_chars])
                else:
                    current_chunk = part
        if current_chunk: chunks.append(current_chunk)
        
        final_segments = []
        current_start = start
        clean_total_len = sum(len(c) for c in chunks)
        
        for chunk in chunks:
            ratio = len(chunk) / clean_total_len if clean_total_len > 0 else 0
            chunk_duration = total_duration * ratio
            chunk_end = current_start + chunk_duration
            final_segments.append({
                "start": current_start,
                "end": chunk_end,
                "text": chunk.strip()
            })
            current_start = chunk_end
        return final_segments

    def process_dual_version(self, raw_segments):
        print(f"   -> 正在进行语义合并 (修复不合理换行)...")
        merged_sentences = []
        buffer_text = ""
        buffer_start = 0.0
        strong_endings = re.compile(r'[。！？\.\!\?]')
        weak_endings = re.compile(r'[，,、]')
        
        for i, seg in enumerate(raw_segments):
            text = seg['text']
            if buffer_text == "": buffer_start = seg['start']
            buffer_text += text
            is_strong = strong_endings.search(text)
            is_long_weak = len(buffer_text) > 150 and weak_endings.search(text)
            is_too_long = len(buffer_text) > 500
            is_end = is_strong or is_long_weak or is_too_long or (i == len(raw_segments)-1)
            
            if is_end:
                full_sent = buffer_text.strip().replace('\n', '').replace('\r', '')
                if full_sent:
                    merged_sentences.append({
                        "start": buffer_start,
                        "end": seg['end'],
                        "text": full_sent
                    })
                buffer_text = ""
        
        print(f"   -> 合并完成。共 {len(merged_sentences)} 个完整句。")
        ai_summary = self.generate_summary(merged_sentences)
        if ai_summary: print("✅ AI 总结生成完毕。")

        final_raw_list = []
        final_ai_list = []
        total = len(merged_sentences)
        
        print("   -> 开始逐句处理 (双轨生成)...")
        for i, item in enumerate(merged_sentences):
            sys.stdout.write(f"\r      进度: {i+1}/{total}...")
            sys.stdout.flush()
            raw_subs = self.smart_split_subtitle(
                item['start'], item['end'], item['text'], MAX_SUBTITLE_LENGTH
            )
            final_raw_list.extend(raw_subs)
            if self.llm_client:
                cleaned_text = self.clean_text_minimal(item['text'])
            else:
                cleaned_text = item['text']
            ai_subs = self.smart_split_subtitle(
                item['start'], item['end'], cleaned_text, MAX_SUBTITLE_LENGTH
            )
            final_ai_list.extend(ai_subs)
            
        print(f"\n   -> 字幕处理完成。")
        return ai_summary, merged_sentences, final_raw_list, final_ai_list

    # ==================================================

    def run(self, user_input, output_task_dir):
        if not self.prepare_source(user_input, output_task_dir): return
        self._load_model()

        print("[3/5] 正在进行语音识别...")
        try:
            result = self.model.transcribe(
                self.video_path, fp16=False, language='zh',
                initial_prompt="以下是简体中文字幕，包含完整的标点符号。"
            )
            ai_summary, merged_sentences, raw_list, ai_list = self.process_dual_version(result["segments"])
        except Exception as e:
            print(f"❌ 处理失败: {e}")
            return

        final_docx_name = f"{self.video_title_stem}.docx"
        final_docx_path = os.path.join(output_task_dir, final_docx_name)
        self.generate_dual_docx(ai_summary, merged_sentences, raw_list, ai_list, final_docx_path)

    # --- 【关键】新增加的字体强制函数 ---
    def add_heading_force_font(self, doc, text, level):
        """添加标题并强制应用中文字体"""
        heading = doc.add_heading(text, level)
        for run in heading.runs:
            run.font.name = u'微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
            # 可选：如果你希望标题是黑色的，而不是Word默认的蓝色，可以取消下面注释
            # run.font.color.rgb = RGBColor(0, 0, 0)
        return heading

    def generate_dual_docx(self, summary, merged_sentences, raw_list, ai_list, output_docx):
        print(f"[4/5] 正在生成 Word 文档...")
        abs_output_path = os.path.abspath(output_docx)
        current_task_dir = os.path.dirname(abs_output_path) 
        self.img_output_dir = os.path.join(current_task_dir, "images")
        if not os.path.exists(self.img_output_dir): os.makedirs(self.img_output_dir)

        doc = Document()
        
        # 基础样式设置 (兜底)
        doc.styles['Normal'].font.name = u'微软雅黑'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')

        # 使用强制字体函数添加标题
        self.add_heading_force_font(doc, self.video_title_stem, 0)
        
        doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"视频路径: {self.video_path}") 
        
        if summary:
            self.add_heading_force_font(doc, '💡 AI 智能总结', 1)
            self.render_markdown_to_word(doc, summary)
            doc.add_page_break()

        self.add_heading_force_font(doc, '第一部分：字幕全文对比', 1)
        
        self.add_heading_force_font(doc, '1.1 原始全文 (未优化)', 2)
        self.add_smart_paragraphs(doc, merged_sentences, use_ai_clean=False)
        
        self.add_heading_force_font(doc, '1.2 AI优化全文 (去口癖/微调)', 2)
        self.add_smart_paragraphs(doc, merged_sentences, use_ai_clean=True)
        
        doc.add_page_break()

        self.add_heading_force_font(doc, '第二部分：图文对照 (AI优化版)', 1)
        doc.add_paragraph("注：此版本去除了口语废话，阅读更流畅。")
        self.create_image_table(doc, ai_list)
        
        doc.add_page_break()
        
        self.add_heading_force_font(doc, '第三部分：图文对照 (原始逐字版)', 1)
        doc.add_paragraph("注：此版本完全忠实于原音频。")
        self.create_image_table(doc, raw_list)

        try:
            doc.save(abs_output_path)
            print(f"\n✅ 完成！\n   📄 文档: {abs_output_path}\n   🖼️ 图片: {self.img_output_dir}\n   🎥视频：{self.video_path}")
        except PermissionError: print(f"❌ 保存失败: 请关闭文档后重试！")
        except Exception as e: print(f"❌ 保存出错: {e}")

    def render_markdown_to_word(self, doc, text):
        for line in text.split('\n'):
            line = line.strip()
            if not line: continue
            if line.startswith('### '):
                self.add_heading_force_font(doc, line.replace('### ', ''), 3)
            elif line.startswith('## '):
                self.add_heading_force_font(doc, line.replace('## ', ''), 2)
            elif line.startswith('**') and line.endswith('**'):
                p = doc.add_paragraph()
                run = p.add_run(line.replace('**', ''))
                run.bold = True
            elif line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif line.startswith('1. '):
                doc.add_paragraph(line, style='List Number')
            else:
                doc.add_paragraph(line)

    def add_smart_paragraphs(self, doc, merged_sentences, use_ai_clean=False):
        current_paragraph_text = ""
        target_length = 500
        for item in merged_sentences:
            if use_ai_clean and self.llm_client:
                text = self.clean_text_minimal(item['text'])
            else:
                text = item['text']
            current_paragraph_text += text
            if len(current_paragraph_text) > target_length:
                p = doc.add_paragraph(current_paragraph_text)
                p.paragraph_format.first_line_indent = Inches(0.3)
                p.paragraph_format.line_spacing = 1.5
                current_paragraph_text = ""
        if current_paragraph_text:
            p = doc.add_paragraph(current_paragraph_text)
            p.paragraph_format.first_line_indent = Inches(0.3)
            p.paragraph_format.line_spacing = 1.5

    def create_image_table(self, doc, segments):
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'
        cap = cv2.VideoCapture(self.video_path)
        
        for i, seg in enumerate(segments):
            start = seg['start']
            text = seg['text'].strip()
            time_str = self.format_time(start).replace(":", "-")
            ms = int((start % 1) * 100) 
            img_filename = f"frame_{time_str}_{ms:02d}.jpg"
            img_path = os.path.join(self.img_output_dir, img_filename)

            if not os.path.exists(img_path) and cap.isOpened():
                mid_point = start + (seg['end']-start)/2 
                cap.set(cv2.CAP_PROP_POS_MSEC, mid_point * 1000)
                ret, frame = cap.read()
                if ret:
                    try: cv2.imwrite(img_path, frame)
                    except: pass
            
            row_cells = table.add_row().cells
            p = row_cells[0].paragraphs[0]
            run_time = p.add_run(f"[{self.format_time(start)}]\n")
            run_time.bold = True
            run_time.font.color.rgb = RGBColor(0, 51, 102) 
            p.add_run(text)
            
            if os.path.exists(img_path):
                try:
                    p_img = row_cells[1].paragraphs[0]
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_img.add_run().add_picture(img_path, width=Inches(3.0))
                except: p_img.add_run("[图片]")

        if cap.isOpened(): cap.release()

    @staticmethod
    def format_time(seconds):
        m, s = divmod(seconds, 60)
        h, m = divmod(m, 60)
        return "%02d:%02d:%02d" % (h, m, s)

# --- 主程序入口 ---
if __name__ == "__main__":
    print("-" * 30)
    print("视频转Word")
    print("-" * 30)
    
    target = input("请输入视频链接 或 本地文件路径:\n>>> ").strip()
    if target:
        base_input = DEFAULT_OUTPUT_BASE
        if not os.path.exists(base_input):
             manual_path = input(f"默认存储路径: {base_input}\n按回车确认，或输入新路径:\n>>> ").strip()
             if manual_path: base_input = manual_path
        
        if not base_input: base_dir = os.getcwd()
        else: base_dir = base_input.strip('"').strip("'")

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        full_task_dir = os.path.join(base_dir, f"Output_{timestamp}")
        
        if not os.path.exists(full_task_dir):
            try: os.makedirs(full_task_dir)
            except: full_task_dir = os.getcwd()

        print(f"\n[设置] 任务目录: {full_task_dir}")
        converter = VideoToWordConverter(model_size="turbo") 
        converter.run(target, full_task_dir)